# Import libraries
import sys
import os
import glob
import re

import pandas as pd
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.shared import Inches, Pt, RGBColor

# Project directory
proj_dir = sys.argv[1]

# Dataset version
version = sys.argv[2]

# Data directory
data_dir  = os.path.join(proj_dir, 'final', 'GRQA_data')

# Metadata directory
meta_dir  = os.path.join(proj_dir, 'final', 'GRQA_meta')

# Figure directory
fig_dir  = os.path.join(proj_dir, 'final', 'GRQA_figures')

# Parameter codes and names
param_stats = pd.read_csv(os.path.join(meta_dir, 'stats', 'GRQA_param_stats.csv'), sep=';')
param_codes = list(param_stats['Parameter code'])
param_names = list(param_stats['Parameter name'])
param_dict = dict(zip(param_codes, param_names))

# Word document object
document = Document()

# Subheadings based on figure name strings
subheadings = {
    'spatial_dist': 'Spatial distribution',
    'temporal_hist': 'Temporal distribution',
    'hist': 'Distribution',
    'box': 'Box plot',
    'availability': 'Monthly time series availability',
    'continuity': 'Monthly time series continuity',
    'median': 'Spatial distribution of yearly median'
}

# Add new heading styles
styles = document.styles

new_heading_style = styles.add_style('New Title', WD_STYLE_TYPE.PARAGRAPH)
new_heading_style.base_style = styles['Title']
font = new_heading_style.font
font.name = 'Arial'
font.size = Pt(20)
font.color.rgb = RGBColor.from_string('000000')
font.bold = True

new_heading_style = styles.add_style('New Heading 1', WD_STYLE_TYPE.PARAGRAPH)
new_heading_style.base_style = styles['Heading 1']
font = new_heading_style.font
font.name = 'Arial'
font.size = Pt(16)
font.color.rgb = RGBColor.from_string('000000')

new_heading_style = styles.add_style('New Heading 2', WD_STYLE_TYPE.PARAGRAPH)
new_heading_style.base_style = styles['Heading 2']
font = new_heading_style.font
font.name = 'Arial'
font.size = Pt(12)
font.color.rgb = RGBColor.from_string('000000')

# Add title
title = f'Data catalog of Global River Water Quality Archive (GRQA) v{version}'
document.add_paragraph(title, style='New Title')
paragraph = document.paragraphs[-1]
paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Figure width
width = Inches(5.9527559055)

# Add all figures to the Word document with corresponding headings
for param_code in param_codes:
    document.add_paragraph('{} ({})'.format(param_dict[param_code], param_code), style='New Heading 1')
    figures = glob.glob(os.path.join(fig_dir, '{}_GRQA*.png'.format(param_code)))
    for fig in figures:
        string = re.search(r'{}_GRQA_(.*?).png'.format(param_code), fig).group(1)
        subheading = '{} of {} observation values'.format(subheadings[string], param_code)
        document.add_paragraph(subheading, style='New Heading 2')
        document.add_picture(fig, width=width)
    paragraph = document.paragraphs[-1]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

# Save as Word document without table of contents (TOC)
fp_docx = os.path.join(
    proj_dir, 'working', f'GRQA_data_catalog_v{version}_without_toc.docx'
)
if os.path.exists(fp_docx):
    os.remove(fp_docx)
document.save(fp_docx)

# TOC has to be added in Word before saving as final PDF
